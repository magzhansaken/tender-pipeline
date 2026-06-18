#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════╗
║  TECHSPEC DUMPER v1.0 — Сбор характеристик лотов               ║
║  Только парсинг, БЕЗ AI, БЕЗ поиска товаров, БЕЗ БД           ║
║  Результат: CSV файл с lot_number|lot_name|tech_spec            ║
╚══════════════════════════════════════════════════════════════════╝

На основе рабочего кода tenderfinder_v2_kaspi.py
Используются те же самые методы парсинга PDF/DOCX/HTML

ЗАПУСК:
  python3 techspec_dumper.py

НАСТРОЙКИ через переменные окружения:
  MAX_PAGES=132       — кол-во страниц (132 страницы × 50 лотов = 6600)
  START_PAGE=1        — с какой страницы начать (для продолжения)
  OUTPUT_FILE=techspecs.csv  — имя выходного файла
"""

import requests
from bs4 import BeautifulSoup
import re
import time
import os
import sys
import csv
import io
from dataclasses import dataclass
from typing import Optional, List, Dict
from datetime import datetime
import random

# Для работы с документами (PDF/DOCX техспецификации)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# ═══════════════════════════════════════════════════════════════════
# НАСТРОЙКИ
# ═══════════════════════════════════════════════════════════════════

MAX_PAGES = int(os.getenv('MAX_PAGES', '132'))       # 132 стр × 50 = 6600 лотов
START_PAGE = int(os.getenv('START_PAGE', '1'))        # С какой страницы начать
OUTPUT_FILE = os.getenv('OUTPUT_FILE', 'techspecs.csv')
DELAY_BETWEEN_LOTS = float(os.getenv('DELAY', '1'))   # Задержка между лотами (сек)
DELAY_BETWEEN_PAGES = float(os.getenv('PAGE_DELAY', '3'))  # Задержка между страницами

print("╔══════════════════════════════════════════════════════════════════╗")
print("║  TECHSPEC DUMPER v1.0 — Сбор характеристик лотов               ║")
print("╚══════════════════════════════════════════════════════════════════╝")
print(f"📦 DOCX парсинг: {'✅' if DOCX_AVAILABLE else '❌ pip install python-docx'}")
print(f"📦 PDF парсинг: {'✅' if PDF_AVAILABLE else '❌ pip install pdfplumber'}")
print(f"📄 Страницы: {START_PAGE} → {MAX_PAGES} ({(MAX_PAGES - START_PAGE + 1) * 50} лотов макс)")
print(f"💾 Выходной файл: {OUTPUT_FILE}")
print(f"⏱️  Задержки: {DELAY_BETWEEN_LOTS}с между лотами, {DELAY_BETWEEN_PAGES}с между страницами")
print()

# ═══════════════════════════════════════════════════════════════════
# DATACLASS
# ═══════════════════════════════════════════════════════════════════

@dataclass
class Lot:
    lot_number: str
    announce_id: str
    lot_url: str
    name: str
    price_per_unit: float
    quantity: int
    total_price: float
    unit: str
    customer: str

# ═══════════════════════════════════════════════════════════════════
# GOSZAKUP PARSER — Только парсинг, без AI
# ═══════════════════════════════════════════════════════════════════

class GoszakupParser:
    """100% из tenderfinder_v2_kaspi.py — только методы парсинга"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        self.base_url = 'https://goszakup.gov.kz'
        self._deadline_cache = {}
        print("✅ Goszakup Parser готов")
    
    def parse_lots_from_search_page(self, page_num: int) -> List[Lot]:
        """100% КОПИЯ из tenderfinder_v2_kaspi.py"""
        url = f'{self.base_url}/ru/search/lots?filter%5Bmethod%5D%5B0%5D=3&filter%5Bstatus%5D%5B0%5D=240&filter%5Bamount_from%5D=150000&filter%5Btrade_type%5D=g&count_record=50&page={page_num}'
        
        try:
            print(f"\n📄 Загружаю страницу {page_num}...")
            response = self.session.get(url, timeout=30)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            rows = soup.select('#search-result tbody tr')
            
            if not rows:
                print(f"   ⚠️  Таблица не найдена или пустая")
                return []
            
            lots = []
            
            for row in rows:
                try:
                    cells = row.find_all('td')
                    if len(cells) < 7:
                        continue
                    
                    # Колонка 0: Номер лота
                    lot_number = cells[0].find('strong')
                    if not lot_number:
                        continue
                    lot_number = lot_number.text.strip()
                    
                    # Колонка 1: Объявление + Заказчик
                    announce_cell = cells[1]
                    announce_link = announce_cell.find('a')
                    if not announce_link:
                        continue
                    
                    announce_url = announce_link.get('href', '')
                    announce_id = announce_url.split('/')[-1] if announce_url else ""
                    
                    # Заказчик
                    customer_elem = announce_cell.find('small', class_='hidden-xs')
                    customer = "Не указан"
                    if customer_elem:
                        customer_text = customer_elem.text
                        if 'Заказчик:' in customer_text:
                            customer = customer_text.split('Заказчик:')[-1].strip()
                    
                    # Колонка 2: Название лота + URL лота
                    lot_cell = cells[2]
                    lot_link = lot_cell.find('a')
                    if not lot_link:
                        continue
                    
                    lot_name = lot_link.find('strong')
                    if not lot_name:
                        continue
                    lot_name = lot_name.text.strip()
                    
                    lot_url = lot_link.get('href', '')
                    if not lot_url.startswith('http'):
                        lot_url = self.base_url + lot_url
                    
                    # Колонка 3: Количество
                    quantity_text = cells[3].text.strip()
                    try:
                        quantity = int(quantity_text.replace(' ', '').replace(',', ''))
                    except:
                        quantity = 1
                    
                    # Колонка 4: Цена (total_price)
                    price_text = cells[4].text.strip()
                    price_match = re.search(r'([\d\s,]+\.\d{2})', price_text)
                    if not price_match:
                        continue
                    
                    total_price = float(price_match.group(1).replace(' ', '').replace(',', ''))
                    price_per_unit = total_price / quantity if quantity > 0 else total_price
                    
                    # Колонка 5: Единица измерения
                    unit = cells[5].text.strip() if len(cells) > 5 else "шт"
                    
                    lot = Lot(
                        lot_number=lot_number,
                        announce_id=announce_id,
                        lot_url=lot_url,
                        name=lot_name,
                        price_per_unit=price_per_unit,
                        quantity=quantity,
                        total_price=total_price,
                        unit=unit,
                        customer=customer
                    )
                    
                    lots.append(lot)
                    
                except Exception as e:
                    continue
            
            print(f"   ✅ Найдено лотов: {len(lots)}")
            return lots
            
        except Exception as e:
            print(f"   ❌ Ошибка загрузки страницы: {e}")
            return []
    
    def get_announce_deadline(self, announce_id: str) -> str:
        """Срок окончания приёма заявок (YYYY-MM-DD HH:MM:SS) со страницы объявления.
        Кэшируется по announce_id — один запрос на объявление, а не на каждый лот."""
        if not announce_id:
            return ""
        if announce_id in self._deadline_cache:
            return self._deadline_cache[announce_id]

        deadline = ""
        try:
            url = f"{self.base_url}/ru/announce/index/{announce_id}"
            resp = self.session.get(url, timeout=15)
            resp.raise_for_status()
            m = re.search(
                r"Срок\s+окончания\s+приема\s+заявок.{0,300}?value=['\"]\s*"
                r"([0-9]{4}-[0-9]{2}-[0-9]{2}(?:[ T][0-9]{2}:[0-9]{2}:[0-9]{2})?)",
                resp.text, re.DOTALL
            )
            if m:
                deadline = m.group(1).strip()
        except Exception:
            deadline = ""

        self._deadline_cache[announce_id] = deadline
        return deadline

    def get_lot_specifications(self, lot_url: str, announce_id: str = None) -> str:
        """Получает техспецификацию ТОЛЬКО из PDF"""
        try:
            extracted_announce_id, lot_id = self._extract_announce_id_from_lot_url(lot_url)
            
            if not announce_id:
                announce_id = extracted_announce_id
            
            if announce_id:
                files = self._get_techspec_files_list(announce_id, lot_id)
                
                if files:
                    for file_info in files:
                        ext = file_info['extension']
                        file_url = file_info['url']
                        
                        # ТОЛЬКО PDF
                        if ext == 'pdf' and PDF_AVAILABLE:
                            text = self._download_and_extract_pdf(file_url)
                            if text:
                                return text[:3000]
            
            return ""
            
        except Exception as e:
            return ""
    
    def _extract_announce_id_from_lot_url(self, lot_url: str) -> tuple:
        """Извлекает announce_id И lot_id из URL лота"""
        match = re.search(r'/subpriceoffer/index/(\d+)/(\d+)', lot_url)
        if match:
            return match.group(1), match.group(2)
        return None, None
    
    def _get_techspec_files_list(self, announce_id: str, lot_id: str = None) -> List[Dict]:
        """Получает список файлов техспецификации"""
        url = f"{self.base_url}/ru/announce/actionAjaxModalShowFiles/{announce_id}/125"
        try:
            response = self.session.get(url, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            all_files = []
            target_file = None
            
            for link in soup.find_all('a', href=True):
                href = link['href']
                if 'v3bl.goszakup.gov.kz/files/download_file/' in href:
                    filename = link.get_text(strip=True)
                    file_info = {
                        'url': href,
                        'filename': filename,
                        'extension': filename.split('.')[-1].lower() if '.' in filename else ''
                    }
                    all_files.append(file_info)
                    
                    if lot_id and lot_id in filename:
                        target_file = file_info
            
            if target_file:
                return [target_file]
            
            return all_files
        except Exception as e:
            return []
    
    def _download_and_extract_docx(self, file_url: str) -> str:
        """Скачивает DOCX и извлекает техническую характеристику"""
        if not DOCX_AVAILABLE:
            return ""
        
        try:
            response = self.session.get(file_url, timeout=30)
            response.raise_for_status()
            
            doc = Document(io.BytesIO(response.content))
            
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            all_text = "\n".join(full_text)
            
            # Ищем поле с техническими характеристиками
            patterns = [
                r'Описание и требуемые\s*функциональные[^:]*характеристики[^:]*товаров[^:]*:\s*(.+?)(?=\n[А-Я]|\nНомер|\nНаименование|\Z)',
                r'функциональные[^:]*технические[^:]*качественные[^:]*эксплуатационные[^:]*характеристики[^:]*:\s*(.+?)(?=\n[А-Я]|\Z)',
                r'характеристики закупаемых товаров:\s*(.+?)(?=\n[А-Я]|\nПриложение|\Z)',
                r'характеристики[^|]*\|\s*(.+?)(?=\n|\|[А-Я]|\Z)',
            ]
            
            for pattern in patterns:
                match = re.search(pattern, all_text, re.DOTALL | re.IGNORECASE)
                if match:
                    spec_text = match.group(1).strip()
                    spec_text = re.sub(r'\s+', ' ', spec_text)
                    if len(spec_text) > 50:
                        return spec_text
            
            # Fallback: поиск по ключевым словам
            keywords = [
                'Описание и требуемые функциональные',
                'технические, качественные и эксплуатационные характеристики',
                'характеристики закупаемых товаров'
            ]
            
            for keyword in keywords:
                if keyword.lower() in all_text.lower():
                    idx = all_text.lower().find(keyword.lower())
                    colon_idx = all_text.find(':', idx)
                    if colon_idx != -1:
                        remaining = all_text[colon_idx + 1:colon_idx + 3000]
                        end_match = re.search(r'\n[А-ЯA-Z][а-яa-z]', remaining)
                        if end_match:
                            remaining = remaining[:end_match.start()]
                        spec_text = re.sub(r'\s+', ' ', remaining).strip()
                        if len(spec_text) > 50:
                            return spec_text
            
            # Если структурированное поле не найдено — весь текст
            if len(all_text) > 100:
                return re.sub(r'\s+', ' ', all_text)[:3000]
            
            return ""
            
        except Exception as e:
            return ""
    
    def _download_and_extract_pdf(self, file_url: str) -> str:
        """Скачивает PDF и извлекает техническую характеристику"""
        if not PDF_AVAILABLE:
            return ""
        
        try:
            response = self.session.get(file_url, timeout=30)
            response.raise_for_status()
            
            with pdfplumber.open(io.BytesIO(response.content)) as pdf:
                
                # Способ 1: Извлекаем таблицы
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if not row or len(row) < 2:
                                continue
                            
                            key = str(row[0] or '').lower()
                            value = str(row[1] or '').strip()
                            
                            if 'характеристики' in key and 'закупаемых товаров' in key:
                                spec_text = re.sub(r'\s+', ' ', value).strip()
                                if len(spec_text) > 50:
                                    return spec_text
                
                # Способ 2: Fallback на текстовый парсинг
                full_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                
                all_text = "\n".join(full_text)
                
                # Ищем русскую секцию
                ru_section = re.search(r'Приложение 2.*', all_text, re.DOTALL | re.IGNORECASE)
                if ru_section:
                    ru_text = ru_section.group(0)
                    
                    match = re.search(r'закупаемых товаров:\s*(.+?)(?=\n[А-Я]|\Z)', 
                                      ru_text, re.DOTALL | re.IGNORECASE)
                    if match:
                        spec_text = match.group(1).strip()
                        spec_text = re.sub(r'\s+', ' ', spec_text).strip()
                        if len(spec_text) > 50:
                            return spec_text
                
                # Fallback: весь текст PDF
                if len(all_text) > 100:
                    return re.sub(r'\s+', ' ', all_text)[:3000]
                
                return ""
                
        except Exception as e:
            return ""
    
    def _get_techspec_from_html_table(self, lot_url: str) -> str:
        """Fallback — парсит таблицу со страницы лота"""
        try:
            response = self.session.get(lot_url, timeout=15)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            
            specs_parts = []
            
            table = soup.find('table', class_='table')
            if table:
                target_fields = [
                    'Наименование ТРУ',
                    'Краткая характеристика', 
                    'Дополнительная характеристика',
                    'Описание'
                ]
                
                for row in table.find_all('tr'):
                    th = row.find('th')
                    td = row.find('td')
                    
                    if th and td:
                        key = th.get_text(strip=True)
                        value = td.get_text(strip=True)
                        
                        for field in target_fields:
                            if field.lower() in key.lower() and value:
                                specs_parts.append(value)
                                break
            
            return "; ".join(specs_parts)
        except Exception as e:
            return ""

# ═══════════════════════════════════════════════════════════════════
# CSV WRITER — Сохранение результатов
# ═══════════════════════════════════════════════════════════════════

class CSVWriter:
    """Пишет результаты в CSV с поддержкой resume"""
    
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.existing_lots = set()
        
        # Читаем уже собранные лоты (для resume)
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter='|')
                    next(reader, None)  # skip header
                    for row in reader:
                        if row:
                            self.existing_lots.add(row[0])
                print(f"📂 Найден существующий файл: {len(self.existing_lots)} лотов уже собрано")
            except:
                self.existing_lots = set()
        
        # Открываем для дозаписи
        file_exists = os.path.exists(filepath) and len(self.existing_lots) > 0
        self.file = open(filepath, 'a', encoding='utf-8', newline='')
        self.writer = csv.writer(self.file, delimiter='|', quoting=csv.QUOTE_ALL)
        
        if not file_exists:
            self.writer.writerow(['lot_number', 'lot_name', 'price_per_unit', 'quantity', 'unit', 'customer', 'deadline', 'tech_spec'])
            self.file.flush()
    
    def is_already_collected(self, lot_number: str) -> bool:
        return lot_number in self.existing_lots
    
    def write_lot(self, lot: Lot, tech_spec: str, deadline: str = ""):
        """Записать один лот"""
        # Очистка tech_spec от переносов строк и pipe символов
        clean_spec = tech_spec.replace('\n', ' ').replace('\r', ' ').replace('|', '/').strip()
        clean_spec = re.sub(r'\s+', ' ', clean_spec)
        
        self.writer.writerow([
            lot.lot_number,
            lot.name,
            f"{lot.price_per_unit:.2f}",
            lot.quantity,
            lot.unit,
            lot.customer,
            deadline,
            clean_spec
        ])
        self.file.flush()
        self.existing_lots.add(lot.lot_number)
    
    def close(self):
        self.file.close()
    
    def count(self):
        return len(self.existing_lots)

# ═══════════════════════════════════════════════════════════════════
# MAIN — Основной цикл сбора
# ═══════════════════════════════════════════════════════════════════

def main():
    print("\n" + "="*70)
    print("🚀 ЗАПУСК СБОРА ХАРАКТЕРИСТИК")
    print("="*70)
    
    parser = GoszakupParser()
    writer = CSVWriter(OUTPUT_FILE)
    
    # Статистика
    stats = {
        'total_lots_seen': 0,
        'skipped_already': 0,
        'collected_with_spec': 0,
        'collected_no_spec': 0,
        'errors': 0,
        'services_skipped': 0
    }
    
    service_keywords = [
        'услуг', 'сервис', 'обслуживание', 'ремонт', 'подписк',
        'монтаж', 'установка', 'наладка', 'консультац', 'аренда',
        'страхован', 'работ по', 'работы по'
    ]
    
    start_time = datetime.now()
    
    for page_num in range(START_PAGE, MAX_PAGES + 1):
        print(f"\n{'='*70}")
        print(f"📄 СТРАНИЦА #{page_num}/{MAX_PAGES}")
        print(f"{'='*70}")
        
        lots = parser.parse_lots_from_search_page(page_num)
        
        if not lots:
            print("   ⚠️  Нет лотов, пропускаю")
            time.sleep(2)
            continue
        
        for idx, lot in enumerate(lots, 1):
            stats['total_lots_seen'] += 1
            
            # Уже собран?
            if writer.is_already_collected(lot.lot_number):
                stats['skipped_already'] += 1
                continue
            
            # НОВОЕ: срок окончания приёма заявок (1 запрос на объявление, с кэшем)
            deadline = parser.get_announce_deadline(lot.announce_id)
            
            # Услуга?
            lot_lower = lot.name.lower()
            is_service = any(kw in lot_lower for kw in service_keywords)
            if is_service:
                stats['services_skipped'] += 1
                # Всё равно записываем, но с пометкой
                writer.write_lot(lot, "[УСЛУГА] " + lot.name, deadline)
                continue
            
            print(f"   📦 [{idx}/{len(lots)}] {lot.lot_number}: {lot.name[:60]}...")
            
            try:
                # Получаем техспецификацию
                tech_spec = parser.get_lot_specifications(lot.lot_url, lot.announce_id)
                
                if tech_spec and len(tech_spec.strip()) > 10:
                    writer.write_lot(lot, tech_spec, deadline)
                    stats['collected_with_spec'] += 1
                    print(f"      ✅ Характеристика: {len(tech_spec)} символов")
                else:
                    # Записываем даже без характеристики (name будет)
                    writer.write_lot(lot, "", deadline)
                    stats['collected_no_spec'] += 1
                    print(f"      ⚠️  Без характеристики")
                
            except Exception as e:
                stats['errors'] += 1
                writer.write_lot(lot, f"[ОШИБКА] {str(e)[:100]}", deadline)
                print(f"      ❌ Ошибка: {e}")
            
            # Задержка между лотами
            time.sleep(DELAY_BETWEEN_LOTS)
        
        # Статистика после каждой страницы
        elapsed = (datetime.now() - start_time).total_seconds()
        total_collected = stats['collected_with_spec'] + stats['collected_no_spec'] + stats['services_skipped']
        speed = total_collected / (elapsed / 60) if elapsed > 0 else 0
        
        print(f"\n   📊 Итого собрано: {writer.count()} лотов")
        print(f"   ✅ С характеристиками: {stats['collected_with_spec']}")
        print(f"   ⚠️  Без характеристик: {stats['collected_no_spec']}")
        print(f"   🏢 Услуги: {stats['services_skipped']}")
        print(f"   ⏭️  Уже были: {stats['skipped_already']}")
        print(f"   ❌ Ошибки: {stats['errors']}")
        print(f"   ⏱️  Скорость: {speed:.0f} лотов/мин | Время: {elapsed/60:.1f} мин")
        
        # Задержка между страницами
        if page_num < MAX_PAGES:
            time.sleep(DELAY_BETWEEN_PAGES)
    
    # Финальная статистика
    writer.close()
    elapsed = (datetime.now() - start_time).total_seconds()
    
    print("\n" + "="*70)
    print("🎉 СБОР ЗАВЕРШЁН!")
    print("="*70)
    print(f"📊 ИТОГО:")
    print(f"   Всего лотов просмотрено: {stats['total_lots_seen']}")
    print(f"   ✅ С характеристиками: {stats['collected_with_spec']}")
    print(f"   ⚠️  Без характеристик: {stats['collected_no_spec']}")
    print(f"   🏢 Услуги (пропущены): {stats['services_skipped']}")
    print(f"   ⏭️  Уже были в файле: {stats['skipped_already']}")
    print(f"   ❌ Ошибки: {stats['errors']}")
    print(f"   📁 Файл: {OUTPUT_FILE} ({writer.count()} записей)")
    print(f"   ⏱️  Общее время: {elapsed/60:.1f} минут")
    print(f"\n💡 Загрузите {OUTPUT_FILE} в Claude для анализа промпта!")

if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️  Остановлено пользователем")
        print(f"💡 Для продолжения: START_PAGE=<номер> python3 techspec_dumper.py")
        sys.exit(0)
    except Exception as e:
        print(f"\n\n❌ КРИТИЧЕСКАЯ ОШИБКА: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
